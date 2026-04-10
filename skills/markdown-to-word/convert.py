"""
Markdown 转 Word 转换工具
将 Markdown 文档转换为 Word 格式
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def convert_markdown_to_word(markdown_content: str, output_path: str = 'output.docx'):
    """
    将 Markdown 内容转换为 Word 文档
    
    Args:
        markdown_content: Markdown 文本内容
        output_path: 输出 Word 文件路径
    
    Returns:
        Document: 创建的 Document 对象
    """
    doc = Document()

    # 配置样式
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 处理内容
    lines = markdown_content.split('\n')
    
    for line in lines:
        # 处理代码块
        if line.startswith('```'):
            # 代码块处理
            continue
        
        # 处理表格
        if is_table(line):
            create_table(doc, line)
        elif is_code_block(line):
            create_code_block(doc, line)
        elif line.startswith('```'):
            # 代码块结束
            create_code_block(doc, '')
        elif is_list_item(line):
            create_list_item(doc, line)
        elif is_header(line):
            create_header(doc, line)
        elif line.strip() == '---':
            # 分隔线
            p = doc.add_paragraph()
            p.add_run('---').italic = True
        else:
            # 普通段落
            p = doc.add_paragraph()
            p.add_run(line.strip())
            # 处理粗体和斜体
            apply_text_formatting(p, line)
    
    # 保存
    doc.save(output_path)
    print(f"文档已保存：{output_path}")
    return doc


def is_header(line: str) -> bool:
    """判断是否为标题"""
    return line.startswith('#')


def is_list_item(line: str) -> bool:
    """判断是否为列表项"""
    return line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('1. ')


def is_code_block(line: str) -> bool:
    """判断是否为代码块"""
    return line.startswith('```')


def is_table(line: str) -> bool:
    """判断是否为表格"""
    return '|' in line


def create_header(doc, line):
    """创建标题"""
    # 提取标题文本
    text = line.lstrip('# ').strip()
    level = line.count('#')
    
    # 设置标题样式
    if level == 1:
        style_name = 'Heading 1'
        font_size = Pt(28)
    elif level == 2:
        style_name = 'Heading 2'
        font_size = Pt(24)
    elif level == 3:
        style_name = 'Heading 3'
        font_size = Pt(20)
    else:
        style_name = 'Heading 1'
        font_size = Pt(18)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    run.font.size = font_size
    run.bold = True
    
    # 应用样式
    for element in p._element.get_or_add_child('w:pPr')[0].get_or_add_child('w:rPr')[0]:
        if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rStyle':
            element.val = style_name


def create_list_item(doc, line):
    """创建列表项"""
    text = line.strip()
    if text.startswith('- '):
        p = doc.add_paragraph(text[2:], style='List Bullet')
    elif text.startswith('* '):
        p = doc.add_paragraph(text[2:], style='List Bullet')
    elif text.startswith('1. '):
        p = doc.add_paragraph(text[2:], style='List Number')
    else:
        p = doc.add_paragraph(text)


def create_code_block(doc, text):
    """创建代码块"""
    p = doc.add_paragraph(style='Code')
    if text:
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.space_after = Pt(12)


def apply_text_formatting(p, text):
    """应用文本格式化"""
    run = p.add_run(text)
    
    # 处理粗体
    if '**' in text:
        # 粗体文本
        parts = text.split('**')
        if len(parts) == 3:
            run.text = parts[0]
            run.bold = False
            new_run = p.add_run(parts[1])
            new_run.bold = True
            new_run.text = parts[2]
    
    # 处理斜体
    if '//' in text:
        # 斜体文本
        parts = text.split('//')
        if len(parts) == 3:
            run.text = parts[0]
            run.italic = False
            new_run = p.add_run(parts[1])
            new_run.italic = True
            new_run.text = parts[2]


def create_table(doc, line):
    """创建表格（简化版）"""
    # TODO: 实现表格解析
    pass


if __name__ == '__main__':
    # 测试示例
    markdown_content = """
# 这是一个测试文档

这是一个普通的段落。

## 二级标题

这里是二级标题的内容。

### 三级标题

这里是三级标题的内容。

- 列表项 1
- 列表项 2
- 列表项 3

**粗体文本** 和 //斜体文本//

这是一个代码块：
```python
print("Hello, World!")
```

表格示例：

| 列 1 | 列 2 |
|------|------|
| 值 1 | 值 2 |
| 值 3 | 值 4 |
    """
    
    convert_markdown_to_word(markdown_content, 'test.docx')